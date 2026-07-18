#!/usr/bin/env python3
"""Convert a finalized brand-research DOCX into a self-contained responsive HTML report."""
from __future__ import annotations
import argparse, html, re
from pathlib import Path
from docx import Document
from docx.document import Document as _Document
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P

URL_RE=re.compile(r'(https?://[^\s<]+)')

def iter_blocks(parent):
    if isinstance(parent,_Document): element=parent.element.body
    elif isinstance(parent,_Cell): element=parent._tc
    else: raise TypeError(type(parent))
    for child in element.iterchildren():
        if isinstance(child,CT_P): yield Paragraph(child,parent)
        elif isinstance(child,CT_Tbl): yield Table(child,parent)

def linkify(text):
    parts=[]; pos=0
    for m in URL_RE.finditer(text):
        parts.append(html.escape(text[pos:m.start()]))
        url=m.group(1).rstrip('.,;)）]')
        trail=m.group(1)[len(url):]
        parts.append(f'<a href="{html.escape(url,quote=True)}" target="_blank" rel="noopener">{html.escape(url)}</a>{html.escape(trail)}')
        pos=m.end()
    parts.append(html.escape(text[pos:]))
    return ''.join(parts).replace('\n','<br>')

def table_html(tbl,index):
    rows=[]
    for ri,row in enumerate(tbl.rows):
        tag='th' if ri==0 else 'td'; cells=[]
        for cell in row.cells:
            text='\n'.join(p.text.strip() for p in cell.paragraphs if p.text.strip())
            cells.append(f'<{tag}>{linkify(text)}</{tag}>')
        rows.append('<tr>'+''.join(cells)+'</tr>')
    return f'<div class="table-wrap" role="region" aria-label="数据表 {index}" tabindex="0"><table>{"".join(rows)}</table></div>'

def main():
    ap=argparse.ArgumentParser()
    ap.add_argument('input_docx'); ap.add_argument('output_html')
    ap.add_argument('--title'); ap.add_argument('--eyebrow',default='BRAND RESEARCH')
    ap.add_argument('--date',default=''); args=ap.parse_args()
    doc=Document(args.input_docx)
    title=args.title or next((p.text.strip() for p in doc.paragraphs if p.style.name=='Title' and p.text.strip()),Path(args.input_docx).stem)
    subtitle=next((p.text.strip() for p in doc.paragraphs if p.style.name=='Subtitle' and p.text.strip()),'公开资料品牌研究')
    blocks=list(iter_blocks(doc)); first_h1=next((i for i,b in enumerate(blocks) if isinstance(b,Paragraph) and b.style.name=='Heading 1'),0)
    intro=[]
    for b in blocks[:first_h1]:
        if isinstance(b,Paragraph) and b.text.strip() and b.style.name not in ('Title','Subtitle') and b.text.strip()!=args.eyebrow:
            intro.append(b.text.strip())
    toc=[]; content=[]; section_open=False; table_i=0; h1_i=0
    for b in blocks[first_h1:]:
        if isinstance(b,Paragraph):
            text=b.text.strip(); style=b.style.name
            if not text: continue
            if style=='Heading 1':
                if section_open: content.append('</section>')
                h1_i+=1; sid=f'section-{h1_i}'; toc.append((sid,text)); section_open=True
                content.append(f'<section class="report-section" id="{sid}"><div class="section-kicker">{h1_i:02d}</div><h2>{linkify(text)}</h2>')
            elif style=='Heading 2': content.append(f'<h3>{linkify(text)}</h3>')
            elif style=='Heading 3': content.append(f'<h4>{linkify(text)}</h4>')
            elif style.startswith('List'):
                content.append(f'<div class="bullet"><span></span><p>{linkify(text)}</p></div>')
            else: content.append(f'<p>{linkify(text)}</p>')
        else:
            table_i+=1; content.append(table_html(b,table_i))
    if section_open: content.append('</section>')
    nav=''.join(f'<a href="#{sid}"><span>{i:02d}</span>{html.escape(text)}</a>' for i,(sid,text) in enumerate(toc,1))
    intro_html=''.join(f'<p>{linkify(x)}</p>' for x in intro)
    css='''
:root{--green:#00754a;--green2:#0b5e40;--ink:#14231d;--muted:#66756e;--paper:#f5f2ea;--card:#fff;--line:#dfe6e1;--mint:#e9f3ee;--amber:#f1b44c;--shadow:0 18px 50px rgba(21,50,38,.10)}
*{box-sizing:border-box}html{scroll-behavior:smooth}body{margin:0;background:var(--paper);color:var(--ink);font-family:Inter,"PingFang SC","Microsoft YaHei",Arial,sans-serif;line-height:1.72}a{color:var(--green2);text-decoration-thickness:1px;text-underline-offset:3px}a:focus-visible{outline:3px solid var(--amber);outline-offset:3px;border-radius:3px}.progress{position:fixed;z-index:99;top:0;left:0;height:4px;background:var(--amber);width:0}.hero{background:radial-gradient(circle at 88% 16%,rgba(241,180,76,.22),transparent 26%),linear-gradient(135deg,#063f2b,#00754a 62%,#149765);color:white;padding:72px max(6vw,32px) 64px}.hero-inner{max-width:1180px;margin:auto}.eyebrow{font-size:12px;letter-spacing:.22em;font-weight:800;opacity:.85}.hero h1{max-width:830px;margin:18px 0 10px;font-size:clamp(38px,6vw,76px);line-height:1.04;letter-spacing:-.04em}.subtitle{font-size:clamp(17px,2vw,25px);max-width:780px;opacity:.9}.hero-meta{display:flex;flex-wrap:wrap;gap:10px;margin-top:30px}.hero-meta span{border:1px solid rgba(255,255,255,.28);background:rgba(255,255,255,.09);padding:7px 12px;border-radius:999px;font-size:13px}.intro{max-width:900px;margin:28px 0 0;padding:22px 24px;border-left:4px solid var(--amber);background:rgba(0,0,0,.12);border-radius:0 12px 12px 0}.intro p{margin:0 0 8px}.intro p:last-child{margin:0}.layout{max-width:1320px;margin:0 auto;display:grid;grid-template-columns:285px minmax(0,1fr);gap:36px;padding:40px 28px 90px}.toc{position:sticky;top:24px;height:calc(100vh - 48px);overflow:auto;background:#0c3d2d;color:white;border-radius:18px;padding:22px 14px;box-shadow:var(--shadow)}.toc-title{padding:0 10px 14px;font-size:12px;font-weight:800;letter-spacing:.15em;opacity:.65}.toc a{display:grid;grid-template-columns:28px 1fr;gap:6px;color:white;text-decoration:none;padding:8px 10px;border-radius:9px;font-size:13px;line-height:1.35}.toc a:hover,.toc a.active{background:rgba(255,255,255,.12)}.toc a span{opacity:.5;font-variant-numeric:tabular-nums}.content{min-width:0}.report-section{scroll-margin-top:20px;background:var(--card);padding:clamp(25px,4vw,48px);margin-bottom:24px;border-radius:20px;box-shadow:var(--shadow);border:1px solid rgba(255,255,255,.9)}.section-kicker{color:var(--green);font-weight:900;font-size:12px;letter-spacing:.18em}.report-section h2{margin:4px 0 22px;font-size:clamp(27px,3vw,42px);line-height:1.16;letter-spacing:-.025em}.report-section h3{margin:32px 0 12px;font-size:22px;color:var(--green2)}.report-section h4{font-size:17px}.report-section>p{max-width:850px}.bullet{display:grid;grid-template-columns:10px 1fr;gap:12px;align-items:start;margin:8px 0}.bullet span{width:7px;height:7px;border-radius:50%;background:var(--amber);margin-top:10px}.bullet p{margin:0}.table-wrap{overflow:auto;margin:24px 0;border:1px solid var(--line);border-radius:14px;background:white}.table-wrap:focus{outline:3px solid rgba(241,180,76,.45)}table{border-collapse:collapse;width:100%;min-width:720px;font-size:13px}th{background:var(--green);color:white;text-align:left;font-size:12px;letter-spacing:.02em;position:sticky;top:0}th,td{padding:12px 13px;border-bottom:1px solid var(--line);vertical-align:top}tbody tr:nth-child(even){background:#f8faf8}tbody tr:hover{background:var(--mint)}td:first-child{font-weight:700;color:#244d3d}.footer{max-width:1180px;margin:0 auto 48px;padding:0 28px;color:var(--muted);font-size:13px}.mobile-nav{display:none}
@media(max-width:900px){.hero{padding:58px 22px 44px}.layout{display:block;padding:20px 14px 60px}.toc{display:none}.mobile-nav{display:block;position:sticky;top:8px;z-index:20;margin-bottom:14px}.mobile-nav select{width:100%;padding:13px 14px;border:1px solid var(--line);border-radius:12px;background:white;color:var(--ink);font-weight:700;box-shadow:0 8px 22px rgba(0,0,0,.08)}.report-section{padding:24px 18px;border-radius:15px}.intro{padding:18px}.table-wrap{margin-left:-4px;margin-right:-4px}}
@media print{body{background:#fff}.progress,.toc,.mobile-nav{display:none}.hero{background:#fff;color:#111;padding:32px 0;border-bottom:3px solid var(--green)}.layout{display:block;padding:20px 0}.report-section{box-shadow:none;border:0;border-bottom:1px solid #ddd;break-inside:auto;padding:18px 0}.table-wrap{overflow:visible}table{font-size:9pt;min-width:0}a{color:#111;text-decoration:none}.footer{padding:0}}
'''
    js='''
const bar=document.querySelector('.progress');const links=[...document.querySelectorAll('.toc a')];const sections=[...document.querySelectorAll('.report-section')];
addEventListener('scroll',()=>{const h=document.documentElement;bar.style.width=(h.scrollTop/(h.scrollHeight-h.clientHeight)*100)+'%';let current=sections[0]?.id;for(const s of sections){if(s.getBoundingClientRect().top<180)current=s.id}links.forEach(a=>a.classList.toggle('active',a.hash==='#'+current))},{passive:true});
document.querySelector('#jump')?.addEventListener('change',e=>document.querySelector(e.target.value)?.scrollIntoView());
'''
    options=''.join(f'<option value="#{sid}">{i:02d} · {html.escape(text)}</option>' for i,(sid,text) in enumerate(toc,1))
    page=f'''<!doctype html><html lang="zh-CN"><head><meta charset="utf-8"><meta name="viewport" content="width=device-width,initial-scale=1"><meta name="color-scheme" content="light"><title>{html.escape(title)}</title><style>{css}</style></head><body><div class="progress" aria-hidden="true"></div><header class="hero"><div class="hero-inner"><div class="eyebrow">{html.escape(args.eyebrow)}</div><h1>{html.escape(title)}</h1><div class="subtitle">{html.escape(subtitle)}</div><div class="hero-meta"><span>公开资料研究</span><span>证据可追溯</span>{f'<span>{html.escape(args.date)}</span>' if args.date else ''}</div><div class="intro">{intro_html}</div></div></header><div class="layout"><aside class="toc" aria-label="报告目录"><div class="toc-title">CONTENTS</div>{nav}</aside><main class="content"><div class="mobile-nav"><label for="jump" class="sr-only">跳转章节</label><select id="jump"><option value="">选择章节…</option>{options}</select></div>{''.join(content)}</main></div><footer class="footer">本报告基于公开资料与定向样本生成；推断与假设需结合甲方内部数据验证。</footer><script>{js}</script></body></html>'''
    Path(args.output_html).write_text(page,encoding='utf-8')
    print(args.output_html)
if __name__=='__main__': main()
